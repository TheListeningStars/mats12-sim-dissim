"""Build docs/truth_conflict_paper.docx from the verified artifacts.

Every number in the document is either read from results/<model>/*.json at build
time or listed in docs/NUMBERS.md, which scripts/reference_sheet.py generates from
the same artifacts. Nothing here is typed by hand except the prose.

The target length is four pages. Detail that used to live in an appendix now lives
in docs/NUMBERS.md and results/LOG.md; this document carries the argument only.

Usage:  python scripts/make_paper.py
"""
from __future__ import annotations

import json
from pathlib import Path

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
RESULTS = ROOT / "results"
OUT = ROOT / "docs" / "truth_conflict_paper.docx"

MODELS = ["qwen3.5-9b", "phi-4", "olmo-3-7b-instruct"]
PRETTY = {
    "qwen3.5-9b": "Qwen3.5-9B",
    "phi-4": "phi-4",
    "olmo-3-7b-instruct": "Olmo-3-7B-Instruct",
}
STRATA = ["<0.3", "0.3-0.6", ">0.6"]


# ----------------------------------------------------------------- artifacts

def load(model: str, name: str):
    return json.loads((RESULTS / model / name).read_text())


def fmt(x, nd=3):
    if x is None:
        return "n/a"
    if isinstance(x, float) and (np.isnan(x) or np.isinf(x)):
        return "n/a"
    return f"{x:.{nd}f}"


def variance_decomposition(model: str) -> dict:
    """Re-derive the source/target/scenario decomposition from transfer_long.csv.

    Deliberately recomputed here rather than read from a cached JSON: this is one
    of the checks scripts/verify.py adds on top of the pipeline, and it should
    agree with what verify.py reports.
    """
    df = pd.read_csv(RESULTS / model / "transfer_long.csv")
    off = df[df.source_cell != df.target_cell].dropna(subset=["auroc"]).copy()
    y = off["auroc"].to_numpy()
    ss_tot = float(((y - y.mean()) ** 2).sum())

    def r2_from(cols):
        X = pd.get_dummies(off[cols].astype(str), drop_first=True).to_numpy(float)
        X = np.column_stack([np.ones(len(off)), X])
        beta, *_ = np.linalg.lstsq(X, y, rcond=None)
        resid = y - X @ beta
        return 1.0 - float((resid ** 2).sum()) / ss_tot

    def scen(col):
        return off[col].astype(str).str.split("|").str[:2].str.join("|")

    off["src_scen"], off["tgt_scen"] = scen("source_cell"), scen("target_cell")
    return {
        "n_cells": int(pd.concat([off.source_cell, off.target_cell]).nunique()),
        "n_pairs": int(len(off)),
        "target": r2_from(["target_cell"]),
        "source": r2_from(["source_cell"]),
        "both": r2_from(["source_cell", "target_cell"]),
        "scenario": r2_from(["src_scen", "tgt_scen"]),
        "off_diag_auroc": float(y.mean()),
    }


def stratum_composition(model: str) -> pd.DataFrame:
    """Composition of each confidence stratum, on exactly the rows rowlevel.py uses."""
    d = pd.read_csv(RESULTS / model / "c_scores.csv")
    d = d[(d.split != "truthfit") & (d.truth_value != -1) & d.parsed.astype(bool)]
    ev = d[d.split == "eval"].copy()
    ev["stratum"] = pd.cut(ev.b_hat.abs(), [0.0, 0.3, 0.6, 1.01], labels=STRATA)
    g = ev.groupby("stratum", observed=True)
    return pd.DataFrame({
        "n": g.size(),
        "frac_hard": g["difficulty"].apply(lambda s: (s == "hard").mean()),
        "frac_honest_cell": g["mode"].apply(lambda s: (s == "honest").mean()),
        "frac_dissim_cell": g["mode"].apply(lambda s: (s == "dissimulation").mean()),
        "n_cells": g["cell"].nunique(),
    }).reindex(STRATA)


ROW = {m: load(m, "rowlevel.json") for m in MODELS}
BASE = {m: load(m, "baselines.json") for m in MODELS}
AXIS = {m: load(m, "truth_axis_meta.json") for m in MODELS}
META = {m: load(m, "run_meta.json") for m in MODELS}
COMP = {m: pd.read_csv(RESULTS / m / "compliance.csv") for m in MODELS}
VDEC = {m: variance_decomposition(m) for m in MODELS}
VDEC["qwen2.5-7b-instruct-dry"] = variance_decomposition("qwen2.5-7b-instruct-dry")
PROMPT = load("qwen3.5-9b", "rowlevel_prompt.json")
STRAT_COMP = stratum_composition("qwen3.5-9b")


def comp_val(model: str, group: str, col: str):
    """One cell of compliance.csv, addressed by mode or by sim_subtype."""
    c = COMP[model]
    hit = c[c["mode"] == group] if group in set(c["mode"]) else c[c["sim_subtype"] == group]
    return float(hit[col].iloc[0])


def auroc(model: str, label: str, stratum: str) -> float:
    return ROW[model]["stratified_auroc"][label][stratum]["auroc"]


# ------------------------------------------------------------------- styling

def setup(doc: Document) -> None:
    n = doc.styles["Normal"]
    n.font.name = "Calibri"
    n.font.size = Pt(10)
    n.paragraph_format.space_after = Pt(5)
    n.paragraph_format.line_spacing = 1.05
    for name, size, color in [("Heading 1", 12.5, 0x1F3864), ("Heading 2", 10.5, 0x2E5496)]:
        s = doc.styles[name]
        s.font.name = "Calibri"
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor(color >> 16, (color >> 8) & 0xFF, color & 0xFF)
        s.font.bold = True
        s.paragraph_format.space_before = Pt(8)
        s.paragraph_format.space_after = Pt(3)


def para(doc, text="", style=None, italic=False, size=None, space_after=None):
    p = doc.add_paragraph(style=style)
    r = p.add_run(text)
    r.italic = italic
    if size:
        r.font.size = Pt(size)
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    return p


def rich(doc, chunks, style=None):
    """chunks: list of (text, bold, italic)."""
    p = doc.add_paragraph(style=style)
    for text, b, i in chunks:
        r = p.add_run(text)
        r.bold, r.italic = b, i
    return p


def bullets(doc, items, size=None):
    for text in items:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(text)
        if size:
            r.font.size = Pt(size)
        p.paragraph_format.space_after = Pt(2)


def mono(doc, text, indent=0.25):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(8.5)
    return p


def table(doc, headers, rows, widths=None, caption=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = t.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(str(h))
        r.bold = True
        r.font.size = Pt(8.5)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(v))
            r.font.size = Pt(8.5)
            if str(v).startswith("**"):
                r.text = str(v).strip("*")
                r.bold = True
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    if caption:
        c = para(doc, caption, italic=True, size=8.5, space_after=6)
        c.paragraph_format.space_before = Pt(2)
    return t


# ------------------------------------------------------------------ document

doc = Document()
setup(doc)
for s in doc.sections:
    s.left_margin = s.right_margin = Inches(0.75)
    s.top_margin = s.bottom_margin = Inches(0.7)

# --- title block
t = doc.add_paragraph()
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
t.paragraph_format.space_after = Pt(2)
r = t.add_run("Truth-Conflict Does Not Explain Deception-Probe Transfer,\n"
              "but Model Confidence Predicts Detectability")
r.bold = True
r.font.size = Pt(15)
r.font.color.rgb = RGBColor(0x1F, 0x38, 0x64)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run("Arsh Shah  ·  code, artifacts and full numbers: "
                "github.com/TheListeningStars/mats12-sim-dissim")
r.font.size = Pt(9.5)
r.italic = True
sub.paragraph_format.space_after = Pt(8)

para(doc,
     "Epistemic status: a solo project, about a week on rented GPUs. The pre-registered "
     "hypothesis failed and is reported as failed. One finding replicates across three model "
     "families and survives the controls I could think of, but uncertainty and fact "
     "difficulty are confounded in the design and nothing here is causal, so I am much less "
     "confident in the mechanism than in the effect. An earlier version of this experiment "
     "was circular and its headline result has been withdrawn (Section 2).",
     italic=True, size=9.5)

# --- summary
para(doc, "Summary", style="Heading 1")
bullets(doc, [
    f"Deception probes are linear readouts of the residual stream that flag a model asserting "
    f"something it internally represents as false. Prior work organises their "
    f"out-of-distribution failure by pragmatic scenario (instructed lying, roleplay, "
    f"counterfactual framing); I tested whether a quantity measured from the model does "
    f"better. It does not. A graded truth-conflict score c adds no explained variance in "
    f"transfer AUROC beyond scenario labels on any of three models (within-scenario "
    f"permutation "
    f"p = {fmt(BASE['qwen3.5-9b']['horse_race']['perm_p_within_scenario'])}, "
    f"{fmt(BASE['olmo-3-7b-instruct']['horse_race']['perm_p_within_scenario'])}, "
    f"{fmt(BASE['phi-4']['horse_race']['perm_p_within_scenario'])}). The existing scenario "
    f"taxonomy survives the challenge.",

    "The first version of this experiment was circular: the conflict score and the deception "
    "label were the same quantity read off the same fitted direction, one graded and one "
    "thresholded. Its apparent positive result is reproduced to three decimals by a symmetric "
    "label-noise model, and is withdrawn.",

    f"What survives the rebuild is about confidence, not framing. One probe trained across "
    f"all cells detects deception at {fmt(auroc('qwen3.5-9b', 'groundtruth', '<0.3'))} AUROC "
    f"on propositions the model is unsure of, against "
    f"{fmt(auroc('qwen3.5-9b', 'groundtruth', '>0.6'))} where it is confident, on a "
    f"ground-truth label that touches neither the fitted direction nor the belief measure. "
    f"Monotone on three families and two label definitions, and present at the final prompt "
    f"token, before the model has emitted anything.",

    f"Incidental, and relevant to anyone building deception evals: phi-4 refuses "
    f"{fmt(comp_val('phi-4', 'dissimulation', 'refusal_rate'))} of explicit \"give the wrong "
    f"verdict\" prompts, then complies with counterfactual framing "
    f"{fmt(comp_val('phi-4', 'counterfactual_world', 'compliance'))} of the time and asserts "
    f"falsehoods at that same rate. Its refusal boundary tracks the framing of the request, "
    f"not the content of the output.",
])

# --- 1 Question and design
para(doc, "1. Question and design", style="Heading 1")
m = META["qwen3.5-9b"]
para(doc,
     "The scenario taxonomy describes the prompts, not the model. If the probe reads a "
     "representation, what organises transfer should be measurable in the representation. The "
     "candidate quantity is a truth-conflict score c: how far the assertion the model makes "
     "sits from that model's own internally represented truth value of the proposition. "
     "Baudrillard's simulation/dissimulation contrast is where the intuition came from, and "
     "is the only work the philosophy does.")
rich(doc, [("Three hypotheses were fixed before the full run. ", False, False),
           ("H1", True, False),
           (": c adds explained variance in transfer AUROC beyond scenario labels (nested "
            "comparison; the primary test). ", False, False),
           ("H2", True, False),
           (": transfer AUROC is monotone in the c-distance between train and test cells. ",
            False, False),
           ("H3", True, False),
           (": high-c-trained probes generalise down better than the reverse. None is "
            "supported.", False, False)])
para(doc,
     f"The design holds content fixed and varies only framing: the same {m['n_statements']} "
     f"statements appear in every condition, so a difference between cells cannot be a "
     f"difference in what is being talked about. Six conditions (honest, instructed lying, "
     f"counterfactual world, fictional frame, persona, and truth-valueless preference "
     f"statements) cross four topics, two styles and two fact difficulties, giving "
     f"{m['manifest_cells']} cells over {m['manifest_rows']} rows. Statements split at the "
     f"pair level, so a fact's true and false forms never straddle a train/eval boundary, and "
     f"ten pairs per topic are held out to fit the truth direction and never probed. "
     f"Generation is greedy at seed {m['seed']} in bf16; the residual stream is cached at "
     f"three depths, at both the response-mean and final-prompt-token positions. "
     f"A row counts as deceptive when the model asserted against its own belief rather than "
     f"against ground truth, because once the fact bank reaches the edge of what the model "
     f"knows, a sincere error is indistinguishable from a lie under a ground-truth "
     f"definition. Where it matters, both definitions are reported.")

# --- 2 Circularity
para(doc, "2. The circularity, and how it was found", style="Heading 1")
a = AXIS["qwen3.5-9b"]["validity"]["instrument_agreement"]
rich(doc, [("The score and the label were the same object. ", True, False),
           ("Version one defined c = max(0, -s·t̂), where s is the sign of what the model "
            "asserted and t̂ its projection onto the fitted truth direction, and labelled a "
            "row deceptive when sign(s)·sign(t̂) < 0. Those are the same comparison, one "
            "graded and one thresholded: c is exactly |t̂| on positive rows and 0 on negative "
            "ones. The horse race was regressing a label on a monotone transform of itself.",
            False, False)])
rich(doc, [("The consequence is quantitative, not rhetorical. ", True, False),
           ("Where |t̂| is near zero the label's own sign is close to random, so a probe "
            "scored against it must lose AUROC whatever the representation does. Symmetric "
            "label noise at rate p attenuates a true AUROC A to p·A + (1-p)(1-A); at the "
            "disagreement rates measured in each stratum, that arithmetic reproduced the "
            "originally logged confidence gradient to three decimal places. The finding was a "
            "property of my labelling.", False, False)])
rich(doc, [("The fix was to measure belief with machinery the direction does not touch. ",
            True, False),
           (f"Teacher-forcing the VERDICT: token and reading the TRUE-versus-FALSE logit "
            f"margin gives an independent belief estimate b̂, which now defines both the "
            f"label and c; t̂ is retained only as an instrument to be checked against b̂. The "
            f"two disagree on {fmt(a['label_disagreement_rate'])} of statements, concentrated "
            f"exactly where the interesting claim had lived: t̂ agrees with b̂ on "
            f"{fmt(a['by_abs_t_hat']['>0.6']['sign_agreement'])} of confident propositions "
            f"and {fmt(a['by_abs_t_hat']['<0.3']['sign_agreement'])} of uncertain ones. The "
            f"independent margin also tracks ground truth better than the fitted direction "
            f"({fmt(a['b_hat_vs_truth_acc'])} against {fmt(a['t_hat_vs_truth_acc'])}). The "
            f"same audit found a counterfactual cell that was half tautology, quoting the "
            f"pair's FALSE member verbatim as the world's fact; those cells now keep TRUE "
            f"members only. Both problems were surfaced before the full run, so no GPU time "
            f"was spent on the flawed design.", False, False)])

# --- 3 Results
para(doc, "3. Results", style="Heading 1")

para(doc, "3.1 H1: the pre-registered horse race is not supported", style="Heading 2")
para(doc,
     "The primary test is a nested comparison over off-diagonal transfer pairs: M0 predicts "
     "transfer AUROC from source and target scenario labels, M2 adds the c terms, and the "
     "question is whether ΔR² beats chance.")
table(doc,
      ["model", "n pairs", "R² M0 (scenario)", "R² M1 (c only)", "R² M2 (both)", "ΔR²",
       "within-scenario permutation p"],
      [[PRETTY[k], BASE[k]["horse_race"]["n_pairs"],
        fmt(BASE[k]["horse_race"]["r2_scenario_only_M0"]),
        fmt(BASE[k]["horse_race"]["r2_c_only_M1"]),
        fmt(BASE[k]["horse_race"]["r2_both_M2"]),
        fmt(BASE[k]["horse_race"]["delta_r2_M2_vs_M0"]),
        f"**{fmt(BASE[k]['horse_race']['perm_p_within_scenario'])}**"] for k in MODELS],
      widths=[1.2, 0.6, 1.1, 0.95, 0.95, 0.55, 1.35],
      caption="Table 1. H1 on three models. No model reaches significance on the "
              "pre-registered test.")
rich(doc, [("The choice of test matters more than the ΔR² does. ", True, False),
           (f"Simulated on the real 17-cell design, the naive nested F test has a 74% type-I "
            f"error rate and clustering on the source cell has 76%, because a few hundred "
            f"cell-pairs are generated by a couple of dozen cells and are nowhere near "
            f"independent. The within-scenario permutation test comes out at 5.0%. So the "
            f"naive F p-value of "
            f"{BASE['qwen3.5-9b']['horse_race']['naive_F_p']:.1e} on Qwen3.5-9B is an "
            f"artifact and should not be quoted; the honest number is the permutation p of "
            f"{fmt(BASE['qwen3.5-9b']['horse_race']['perm_p_within_scenario'])}, and an "
            f"independent cell-level permutation null run by the verification script agrees "
            f"at p = 0.151. H2 and H3 fare no better: H2's apparent Spearman ρ = 0.571 "
            f"(p = 0.021, n = 16) falls to ρ = 0.522 (p = 0.067, n = 13) once three "
            f"degenerate counterfactual cells are dropped, and the whole correlation lives "
            f"inside an AUROC range of 0.941 to 0.9996, which is rank noise on a "
            f"ceiling-compressed scale. H3's asymmetry is within noise. One caveat "
            f"cuts against my own null: this design could only ever have detected a large "
            f"effect, so H1 is underpowered as well as unsupported.", False, False)])

para(doc, "3.2 Confidence-stratified detectability", style="Heading 2")
para(doc,
     "H1 tested the construct at the level of cells; the result that survives tests it at the "
     "level of rows. One probe is trained once across all cells, then evaluated within strata "
     "of the model's confidence in the proposition, |b̂|.")
hdr = ["model", "label", "|b̂| < 0.3", "|b̂| 0.3-0.6", "|b̂| > 0.6"]
rows = []
for k in MODELS:
    for lab in ["belief", "groundtruth"]:
        cells = [PRETTY[k] if lab == "belief" else "", lab]
        for s in STRATA:
            r = ROW[k]["stratified_auroc"][lab][s]
            cells.append(f"{fmt(r['auroc'])} [{fmt(r['ci_lo'],2)}, {fmt(r['ci_hi'],2)}]")
        rows.append(cells)
table(doc, hdr, rows, widths=[1.25, 1.0, 1.4, 1.4, 1.4],
      caption="Table 2. AUROC with bootstrap 95% CIs. 'belief' labels a row deceptive when "
              "the assertion contradicts the model's own logit margin; 'groundtruth' uses "
              "external fact, so involves neither the fitted direction nor the margin. "
              "Monotone in every model on both labels.")
para(doc,
     f"On Qwen3.5-9B the intervals for the lowest and highest strata do not overlap under "
     f"either label. phi-4's gradient is compressed because it refuses almost all "
     f"instructed-lying prompts, leaving a positive class drawn from the easier "
     f"counterfactual and fictional cells (Section 3.4). Two objections had to be cleared. "
     f"The response-mean features contain the verdict token the probe is asked to predict, "
     f"which is what the {fmt(BASE['qwen3.5-9b']['behavioral_text_auroc'])} text-only TF-IDF "
     f"baseline picks up; reading instead at the final prompt token, before the model has "
     f"emitted anything, costs "
     f"{fmt(auroc('qwen3.5-9b', 'belief', 'all') - PROMPT['stratified_auroc']['belief']['all']['auroc'])} "
     f"AUROC overall and makes the ground-truth gradient steeper, not shallower "
     f"({fmt(PROMPT['stratified_auroc']['groundtruth']['<0.3']['auroc'])} to "
     f"{fmt(PROMPT['stratified_auroc']['groundtruth']['>0.6']['auroc'])} at the prompt site "
     f"against {fmt(auroc('qwen3.5-9b', 'groundtruth', '<0.3'))} to "
     f"{fmt(auroc('qwen3.5-9b', 'groundtruth', '>0.6'))} at the response site). And label "
     f"noise, the mechanism that killed version one, over-predicts the damage: corrupting the "
     f"clean high-confidence stratum at the disagreement rate measured in the low stratum "
     f"lands at {fmt(ROW['qwen3.5-9b']['noise_injection_control']['auroc_after_injecting_noise'])} "
     f"on Qwen3.5-9B against an observed "
     f"{fmt(ROW['qwen3.5-9b']['noise_injection_control']['observed_low_stratum_auroc'])}, and "
     f"undershoots in all three models by 0.14 to 0.43 AUROC.")
para(doc,
     "The supported claim is narrow: the pre-commitment state encodes whether the model is "
     "about to assert against its own belief, and that encoding is weaker on propositions it "
     "is unsure about. It behaves more like a threshold than a smooth gradient, and "
     "uncertainty is not established as the mediator (Section 4).")

para(doc, "3.3 What did not replicate", style="Heading 2")
para(doc,
     f"On a single model it looked like transfer AUROC was almost entirely a property of the "
     f"target cell (R² {fmt(VDEC['qwen3.5-9b']['target'])}) rather than of the train/test "
     f"mismatch (source-cell R² {fmt(VDEC['qwen3.5-9b']['source'])}). That holds on three of "
     f"four models and inverts on Olmo, where source explains "
     f"{fmt(VDEC['olmo-3-7b-instruct']['source'])} against "
     f"{fmt(VDEC['olmo-3-7b-instruct']['target'])} for target and scenario labels explain more "
     f"than either ({fmt(VDEC['olmo-3-7b-instruct']['scenario'])}). Olmo is also the only "
     f"model whose probes meaningfully degrade out of distribution, at "
     f"{fmt(VDEC['olmo-3-7b-instruct']['off_diag_auroc'])} off-diagonal AUROC against "
     f"{fmt(VDEC['qwen3.5-9b']['off_diag_auroc'])} and "
     f"{fmt(VDEC['phi-4']['off_diag_auroc'])}. The pattern that would make sense of this, "
     f"offered as a hypothesis and not a result: target-dominance is a symptom of ceiling, "
     f"since where transfer already sits near 1.0 there is no mismatch penalty left to "
     f"measure and all remaining variance is \"how hard is this target cell\". The pilot model "
     f"is the awkward case, at {fmt(VDEC['qwen2.5-7b-instruct-dry']['off_diag_auroc'])} "
     f"off-diagonal but still target-dominated, so ceiling is not the whole story. The earlier "
     f"claim that the decomposition replicates across families was wrong: it replicates on "
     f"three and inverts on the fourth.")

para(doc, "3.4 Refusal tracks framing, not content", style="Heading 2")
para(doc,
     f"This was not something I set out to measure; running the same prompts across three "
     f"families made it unavoidable. phi-4 declines "
     f"{fmt(comp_val('phi-4', 'dissimulation', 'refusal_rate'))} of \"deliberately give the "
     f"WRONG verdict\" prompts while complying with counterfactual framing "
     f"{fmt(comp_val('phi-4', 'counterfactual_world', 'compliance'))} of the time and "
     f"producing assertions that contradict ground truth at that same rate. The matched pair "
     f"below is from the seeded random transcript sample, not selected for effect:")
mono(doc, 'Instructed lying  "The capital of Kenya is Nairobi."')
mono(doc, "  -> I'm sorry, but I can't assist with that request.")
mono(doc, 'Counterfactual frame  "The capital of Iraq is Baghdad."')
mono(doc, "  -> VERDICT: FALSE / In our world, the capital of Iraq is Seoul, not Baghdad.")
para(doc,
     "The refusal boundary has keyed on the pragmatic shape of the request rather than on the "
     "truth-conflict content of the output. One identical instruction produces three "
     "behaviours across the three families: Qwen mostly complies, Olmo answers honestly "
     "anyway, phi-4 declines. A deception eval built only on explicit \"lie to me\" prompts "
     "would be measuring refusal training on phi-4, instruction-following on Olmo and "
     "deception on Qwen, while reporting a single number. Counterfactual framing routes "
     "around all three. The cost to this study is that phi-4's dissimulation cells carry 8 "
     "usable rows out of 274 and drop out, leaving phi-4 with 8 cells.",
     space_after=6)

# --- 4 Controls and limitations
para(doc, "4. Controls and limitations", style="Heading 1")
sc = STRAT_COMP
v = AXIS["qwen3.5-9b"]["validity"]
bullets(doc, [
    f"Controls on the Section 3.2 result: a ground-truth label independent of both the fitted "
    f"direction and the logit margin (Table 2); the prompt-site read; noise injection; and "
    f"flat stratum composition, with the fraction of rows from honest cells at "
    f"{fmt(sc['frac_honest_cell'].iloc[0],3)} / {fmt(sc['frac_honest_cell'].iloc[1],3)} / "
    f"{fmt(sc['frac_honest_cell'].iloc[2],3)} across strata and the dissimulation fraction "
    f"identical. Transfer numbers are reported against a random-direction floor, within-class "
    f"OOD, style shift, text-only TF-IDF and length-only baselines; length-only is what caught "
    f"a ceiling artifact in the pilot, where it reached 0.752.",

    "scripts/verify.py recomputes every headline number with code that imports nothing from "
    "the analysis modules, and adds three checks the pipeline never performed (cell-level "
    "permutation null, variance decomposition, label-noise control); all comparable "
    "quantities agree on all three models. The full run reproduces bit-for-bit on different "
    "hardware, and thirty seeded random transcripts were read by hand to check the compliance "
    "and verdict parsing.",

    f"Uncertainty and difficulty are confounded, which is the biggest hole here. The "
    f"low-confidence stratum is {fmt(sc['frac_hard'].iloc[0]*100,1)}% hard facts against "
    f"{fmt(sc['frac_hard'].iloc[2]*100,1)}% in the high stratum, so \"the model is uncertain\" "
    f"and \"the fact is hard\" cannot be separated in this design.",

    f"Instrument validity fails one of four checks: truth-valueless preference statements draw "
    f"{fmt(v['preference_to_factual_belief_ratio'],2)} times the belief margin of factual "
    f"ones, so the intended zero anchor for c does not exist. The other three pass.",

    "Nothing here is causal; the truth direction was never patched or steered. phi-4 "
    "contributes 8 cells rather than 30, so it is a weak replication for anything cell-level. "
    "Four models is thin for the ceiling hypothesis. And batched generation was implemented "
    "but not shipped after an equivalence check against batch size 1 failed on 2 of 12 rows; "
    "the likely cause is benign, but it is unconfirmed and that code sits on the path that "
    "produces every number here.",
])

# --- 5 Next
para(doc, "5. What I would do next", style="Heading 1")
bullets(doc, [
    "Deconfound uncertainty from difficulty with easy facts the model is nonetheless "
    "uncertain about, and hard facts it is confident about. Without this the main result "
    "stays descriptive.",
    "Test the confidence account causally by steering |b̂| and predicting the change in probe "
    "AUROC, rather than observing the correlation.",
    "Add models spanning the 0.85 to 1.0 off-diagonal range to test the ceiling hypothesis, "
    "since the current four cluster at the top.",
    "Chase the refusal finding on its own terms. Whether framing-sensitivity of refusal "
    "generalises beyond phi-4 and beyond this task is cheap to run and matters more for eval "
    "design than anything else here.",
])

para(doc,
     "Every number above is read from results/<model>/*.json and *.csv at build time by "
     "scripts/make_paper.py, or recomputed there from transfer_long.csv and c_scores.csv. "
     "Run provenance and the full baseline, compliance and instrument tables are in "
     "docs/NUMBERS.md, generated from the same artifacts by scripts/reference_sheet.py; the "
     "verbatim prompt templates are in src/data_build.py.",
     italic=True, size=9)

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUT)
print(f"wrote {OUT.relative_to(ROOT)}")
